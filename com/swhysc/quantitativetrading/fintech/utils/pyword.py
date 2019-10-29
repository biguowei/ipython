# coding:utf-8

"""
    python 操作word
"""

from docx import Document
from docx.shared import Inches

# 创建word文档对象
document = Document()
# 添加标题
document.add_heading('Document Title', 1)

# 添加段落
p = document.add_paragraph('A plain paragraph having some ')

# 添加段落文字，并指定样式：这里设置文字为加粗
p.add_run('bold').bold = True

# 添加段落文字
p.add_run(' and some ')

# 添加段落文字，并指定样式：设置文字为斜体
p.add_run('italic.').italic = True

# 添加一级标题
document.add_heading('Heading, level 1', level=1)

# 添加段落，并设置段落样式
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'  # 样式为一个小圆点
)
document.add_paragraph(
    'first item in ordered list', style='List Number'  # 样式为数字
)

# 插入图片:Inches表示以英寸作为图片的单位
document.add_picture('pic.jpg', width=Inches(3.0))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# 创建一行三列的表格
table = document.add_table(rows=1, cols=3)
# 获取第一行的所有列数
hdr_cells = table.rows[0].cells

# 给第一行的各个列添加内容
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

# 给table表格添加新行，并给各列添加内容
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

# 添加分页
# document.add_page_break()

# 保存world文档
document.save('demo.docx')

if __name__ == "__main__":
    pass