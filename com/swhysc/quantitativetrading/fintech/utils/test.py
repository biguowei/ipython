import re
import sys
import time
from collections import defaultdict
from imp import reload

import easygui as g
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QFileDialog, QWidget
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook


class MyWindow(QWidget):
    def __init__(self):
        super(MyWindow, self).__init__()
        file_path, filetype = QFileDialog.getOpenFileName(self, "选择文件", "/", "All Files (*);;Text Files (*.txt)")
        print(file_path)  # 打印文件全部路径（包括文件名和后缀名）
        # 获取文件名
        file_name = re.findall(r'[^\\/:*?"<>|\r\n]+$', file_path)
        file_name = re.findall(r'(.+?)\.xlsx', file_name[0])
        print(file_name[0] + '.xlsx')
        try:
            reload(sys)
            # 开始时间
            startTime = time.time()

            # 读取excel xlsx文件
            wb = load_workbook(file_path)

            # 获取所有sheet页名字
            xl_sheet_names = wb.get_sheet_names()

            # 定位到相应sheet页,[0]为sheet页索引
            xl_sheet = wb.get_sheet_by_name(xl_sheet_names[0])

            # 获取行列数
            excel_row = xl_sheet.max_row
            excel_column = xl_sheet.max_column

            # 取excel第一页第一张表
            i = 0

            # 将excel表格装入itercars操作

            itercars = iter(xl_sheet.rows)
            iters = iter(xl_sheet.rows)
            """
            for row in xl_sheet.rows:
                for num in range(0, excel_column):
                    print(u'' + str(row[num].value))
                break
            """
            # 进入第二行开始循环插入
            next(itercars)
            next(iters)

            # 读取Excel中内容，以学校为key生成字典{"graduate_institutions":{"name","cert_no","major","education","award_year","award_detail"},...}
            results = defaultdict(list)

            # 先拿到所有的学校,将所有的数据放入数据字典
            """
            graduate_institutions = []

            for row in iters:
                graduate_institutions.append(row[2].value)
                results[row[2].value].append([])
            """

            for row in itercars:
                content = [row[2].value, row[0].value, row[1].value, row[4].value, row[3].value, row[5].value,
                           row[6].value]
                results[row[2].value].append(content)

            # 以学校为文件名输出各高校的核查文件
            # print(results)
            for item in list(results.items()):
                # 逐个学校写word文件
                # print(item[1])
                # 创建word文档对象
                document = Document()
                p = document.add_paragraph('附件：核查名单')

                # 创建一行九列的表格
                table = document.add_table(rows=1, cols=9, style="Table Grid")
                # 获取第一行的所有列数
                hdr_cells = table.rows[0].cells

                # 给第一行的各个列添加内容
                hdr_cells[0].text = '序号'
                hdr_cells[1].text = '学校'
                hdr_cells[2].text = '姓名'
                hdr_cells[3].text = '身份证号'
                hdr_cells[4].text = '学历'
                hdr_cells[5].text = '专业'
                hdr_cells[6].text = '获奖年份'
                hdr_cells[7].text = '获奖名称'
                hdr_cells[8].text = '核查结果反馈'

                # 给table表格添加新行，并给各列添加内容
                # print(item[1])

                idx = 0
                for graduate_institution, name, cert_no, education, major, award_year, award_detail in item[1]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = str(graduate_institution)
                    row_cells[2].text = str(name)
                    row_cells[3].text = str(cert_no)
                    row_cells[4].text = str(education)
                    row_cells[5].text = str(major)
                    row_cells[6].text = str(award_year)
                    row_cells[7].text = str(award_detail)
                    row_cells[8].text = ""
                    idx += 1

                p = document.add_paragraph('经办人（签名）：                                    联系电话：             ')
                p = document.add_paragraph('                                               学校盖章：')
                # 设置段落水平右对齐
                paragraph_format = p.paragraph_format
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p = document.add_paragraph('                                               日    期：')
                # 设置段落水平右对齐
                paragraph_format = p.paragraph_format
                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p = document.add_paragraph('备注：1.核查结果反馈为：“属实”或“不属实”，如有特殊情况可详细说明；')

                # table.style.font.name = u'黑体'

                # 保存world文档
                document.save(str(item[1][0][0]) + '.docx')

        except Exception as e:
            print(e)
            title = g.msgbox(msg="               生成失败:" + e, title="Error", ok_button="确定")


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    myshow = MyWindow()
    myshow.show()
    myshow.close()
    sys.exit()
