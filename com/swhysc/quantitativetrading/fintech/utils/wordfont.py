#!/usr/bin/python3
# -*- coding: utf-8 -*-
# @Time    : 2019/2/25 9:52
# @File    : table_add_cell.py


from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import os

"""
word表格,添加一个单元格以及内容
"""
path = "F:\\format.docx"  # 文件路径
document = Document(os.path.abspath(path))  # 读入文件
# 设置字体为： 宋体
document.styles['Normal'].font.name = u'宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

tables = document.tables  # 获取文件中的tables
print(len(tables))
# 被修改表格的列表
one_cells = []


def match_table():
    """
    遍历文档匹配满足条件的表格
    :return:
    """
    for table in tables:
        rows_num = len(table.rows)
        columns_num = len(table.columns)
        if rows_num > 2 and columns_num == 9:
            one_text = table.cell(0, 1).text
            if len(one_text) == 29:
                one_cells.append(one_text)

                end_text = table.cell(rows_num - 1, columns_num - 1).text
                if end_text:
                    # 添加一行
                    table.add_row()
                    set_font(table, 5, 0)
                else:
                    set_font(table, rows_num - 1, columns_num - 2)
    document.save("F:\\format.docx")


def set_font(table, a, b):
    """
    设置表格字体样式
    :param table: 表格
    :param a: 行坐标
    :param b: 列坐标
    :return:
    """
    run = table.cell(a, b).paragraphs[0].add_run(u'信息资源生产格式')  # 内容为：信息资源生产格式
    run.bold = True  # 加粗
    run.font.size = Pt(9)  # 字体大小：小五-9
    table.cell(a, b).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    run1 = table.cell(a, b + 1).paragraphs[0].add_run(u'ORACLE')
    run1.font.size = Pt(9)
    table.cell(a, b + 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def check_data():
    # 校验修改内容
    for i in one_cells:
        print(i)


if __name__ == '__main__':
    print(1)

    match_table()
    check_data()

    print(2)